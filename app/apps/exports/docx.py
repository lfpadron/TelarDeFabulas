from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from django.utils.translation import gettext as _

from apps.manuscripts.models import ManuscriptNode

from .services import ordered_export_nodes


HEADING_LEVELS = {
    ManuscriptNode.NodeType.BOOK: 1,
    ManuscriptNode.NodeType.PART: 2,
    ManuscriptNode.NodeType.CHAPTER: 2,
    ManuscriptNode.NodeType.SCENE: 3,
    ManuscriptNode.NodeType.FRAGMENT: 4,
}

ALIGNMENTS = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
}


def as_float(value):
    return float(value)


def split_blocks(content):
    return [block.strip() for block in (content or "").split("\n\n") if block.strip()]


def apply_document_styles(document, style_template):
    normal = document.styles["Normal"]
    normal.font.name = style_template.font_body
    normal.font.size = Pt(as_float(style_template.body_size))
    normal.paragraph_format.line_spacing = as_float(style_template.line_spacing)
    normal.paragraph_format.space_after = Pt(as_float(style_template.paragraph_spacing))
    normal.paragraph_format.first_line_indent = Mm(as_float(style_template.first_line_indent))
    normal.paragraph_format.alignment = ALIGNMENTS.get(style_template.text_alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    for level in range(1, 5):
        heading = document.styles[f"Heading {level}"]
        heading.font.name = style_template.font_heading
        heading.font.size = Pt(as_float(style_template.heading_size))
        heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.space_after = Pt(as_float(style_template.paragraph_spacing))

    for section in document.sections:
        section.top_margin = Mm(as_float(style_template.margin_top))
        section.bottom_margin = Mm(as_float(style_template.margin_bottom))
        section.left_margin = Mm(as_float(style_template.margin_left))
        section.right_margin = Mm(as_float(style_template.margin_right))


def apply_normal_paragraph_format(paragraph, style_template):
    paragraph.alignment = ALIGNMENTS.get(style_template.text_alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
    paragraph.paragraph_format.line_spacing = as_float(style_template.line_spacing)
    paragraph.paragraph_format.space_after = Pt(as_float(style_template.paragraph_spacing))
    paragraph.paragraph_format.first_line_indent = Mm(as_float(style_template.first_line_indent))
    for run in paragraph.runs:
        run.font.name = style_template.font_body
        run.font.size = Pt(as_float(style_template.body_size))


def add_table_of_contents(document, nodes):
    if not nodes:
        return

    document.add_heading(_("Índice"), level=1)
    for node in nodes:
        level = HEADING_LEVELS.get(node.node_type, 2)
        indent = Mm((level - 1) * 6)
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.left_indent = indent
        paragraph.paragraph_format.first_line_indent = Mm(0)
        paragraph.add_run(node.title)
    document.add_page_break()


def add_node(document, node, style_template):
    heading_level = HEADING_LEVELS.get(node.node_type, 2)
    heading = document.add_heading(node.title, level=heading_level)
    for run in heading.runs:
        run.font.name = style_template.font_heading
        run.font.size = Pt(as_float(style_template.heading_size))

    for block in split_blocks(node.content):
        paragraph = document.add_paragraph(style="Normal")
        paragraph.add_run(block.replace("\n", "\n"))
        apply_normal_paragraph_format(paragraph, style_template)

    if node.node_type == ManuscriptNode.NodeType.SCENE and style_template.scene_separator:
        separator = document.add_paragraph(style="Normal")
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator.paragraph_format.first_line_indent = Mm(0)
        separator.add_run(style_template.scene_separator)


def build_export_docx(export_job):
    document = Document()
    style_template = export_job.style_template
    nodes = ordered_export_nodes(export_job)
    apply_document_styles(document, style_template)

    if style_template.include_table_of_contents:
        add_table_of_contents(document, nodes)

    for node in nodes:
        add_node(document, node, style_template)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
